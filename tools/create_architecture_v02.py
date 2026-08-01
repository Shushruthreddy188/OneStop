from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION

OUT = r"C:\Users\ShushruthKumarReddy\OneDrive\Desktop\OneStop\OneStop_MVP_Architecture_v0.2.docx"

NAVY = "17324D"
BLUE = "2374AB"
TEAL = "2A9D8F"
PALE = "EAF2F8"
GRAY = "5B6573"
WHITE = "FFFFFF"

doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(.65)
sec.bottom_margin = Inches(.65)
sec.left_margin = Inches(.72)
sec.right_margin = Inches(.72)

styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(9.5)
styles["Normal"].font.color.rgb = RGBColor.from_string("263442")
styles["Normal"].paragraph_format.space_after = Pt(5)
for name, size, color in [("Title", 30, NAVY), ("Heading 1", 18, NAVY), ("Heading 2", 12, BLUE)]:
    styles[name].font.name = "Aptos Display"
    styles[name].font.size = Pt(size)
    styles[name].font.color.rgb = RGBColor.from_string(color)
    styles[name].font.bold = True

def shade(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), color); tcPr.append(shd)

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr(); el = OxmlElement('w:tblHeader'); el.set(qn('w:val'), 'true'); trPr.append(el)

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Light Shading Accent 1"
    for i, h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h; shade(c, NAVY); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for r in c.paragraphs[0].runs: r.font.bold=True; r.font.color.rgb=RGBColor.from_string(WHITE); r.font.size=Pt(8.5)
    set_repeat_table_header(t.rows[0])
    for ridx, row in enumerate(rows):
        cells=t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text=str(v); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after=Pt(2)
                for r in p.runs: r.font.size=Pt(8.3)
        if ridx % 2 == 0:
            for c in cells: shade(c, "F4F7FA")
    return t

def bullet(text):
    p=doc.add_paragraph(style="List Bullet"); p.add_run(text); return p

def page_break(): doc.add_page_break()

# Cover
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(70)
r=p.add_run("ONESTOP"); r.bold=True; r.font.size=Pt(15); r.font.color.rgb=RGBColor.from_string(TEAL)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("MVP Architecture"); r.bold=True; r.font.size=Pt(34); r.font.color.rgb=RGBColor.from_string(NAVY)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Implementation baseline & release-hardening plan"); r.font.size=Pt(15); r.font.color.rgb=RGBColor.from_string(GRAY)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(24)
r=p.add_run("Version 0.2  |  31 July 2026  |  Status: MVP feature path complete"); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=RGBColor.from_string(BLUE)

box=doc.add_table(rows=1, cols=1); box.alignment=WD_TABLE_ALIGNMENT.CENTER; box.autofit=False; box.columns[0].width=Inches(5.8)
c=box.cell(0,0); shade(c, PALE)
p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(12); p.paragraph_format.space_after=Pt(12)
r=p.add_run("Architecture intent\nA small, understandable retail platform whose checkout remains correct across retries, concurrent requests, and temporary messaging failures."); r.font.size=Pt(11)

page_break()
doc.add_heading("1. Executive summary", level=1)
doc.add_paragraph("OneStop now implements the full MVP customer journey—from authentication and catalog discovery through cart, inventory-backed checkout, order history, cancellation, and asynchronous confirmation notifications. The system remains a microservice architecture, but correctness is concentrated at the checkout boundary where failures have business impact.")
doc.add_heading("MVP assessment", level=2)
table(["Area", "Status", "Evidence"], [
    ("Customer journey", "Complete", "Register/login, browse/search, cart, checkout, history, cancellation"),
    ("Inventory correctness", "Complete", "Atomic reservation, confirmation, release and cancellation restock"),
    ("Retry safety", "Complete", "Customer-scoped idempotency key and PostgreSQL unique constraint"),
    ("Interrupted checkout", "Complete", "Durable STOCK_RESERVED handoff plus reconciliation worker"),
    ("Notification delivery", "Complete for MVP", "Transactional outbox, Kafka retry, idempotent consumer"),
    ("Production readiness", "In progress", "CI complete; tracing, alerts, secrets and contract tests remain"),
])
doc.add_heading("What changed since v0.1", level=2)
for x in [
    "All planned domain services and the React customer application are implemented and integrated.",
    "Checkout now persists PENDING → STOCK_RESERVED → CONFIRMED with a recoverable inventory correlation id.",
    "Order confirmation and its outbox record commit atomically; Kafka delivery retries with exponential backoff.",
    "PostgreSQL/Testcontainers verifies concurrent checkout idempotency, and Docker smoke testing verifies the complete path.",
]: bullet(x)

doc.add_heading("2. Scope and boundaries", level=1)
doc.add_paragraph("The MVP optimizes for a dependable demonstration and a clean foundation, not marketplace-scale breadth.")
table(["In scope", "Deferred"], [
    ("JWT customer identity and profile", "Social login, MFA, password recovery"),
    ("Catalog browse/search and product detail", "Recommendations, reviews, advanced merchandising"),
    ("Persistent cart", "Promotions, coupons, saved lists"),
    ("Inventory reservation and stock restoration", "Warehouses, sourcing and fulfillment routing"),
    ("COD/card selection without payment capture", "Payment gateway, refunds and settlement"),
    ("Order confirmation notification log", "Production email/SMS provider and templates"),
])

page_break()
doc.add_heading("3. System architecture", level=1)
doc.add_paragraph("The browser talks only to the API Gateway. Each backend service owns its data and publishes a bounded API. Kafka is used for order-confirmation events; synchronous HTTP remains appropriate for the checkout commands that require an immediate business decision.")

arch = [
    "React SPA (5173)",
    "        │ HTTPS / JSON",
    "API Gateway (8080)",
    "   ├── Identity (8081) ───── PostgreSQL identity_db",
    "   ├── Catalog (8082) ────── PostgreSQL catalog_db",
    "   ├── Cart (8083) ───────── PostgreSQL cart_db",
    "   ├── Inventory (8084) ──── PostgreSQL inventory_db",
    "   └── Order (8085) ──────── PostgreSQL order_db",
    "          │ transactional outbox",
    "          └── Kafka: order.confirmed ── Notification (8086) ── PostgreSQL notification_db",
]
p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(.25); p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(12)
r=p.add_run("\n".join(arch)); r.font.name="Cascadia Mono"; r.font.size=Pt(9); r.font.color.rgb=RGBColor.from_string(NAVY)

doc.add_heading("Service responsibilities", level=2)
table(["Component", "Owns", "Key responsibility"], [
    ("Gateway", "Routing", "Single entry point, CORS and service routing"),
    ("Identity", "Users, addresses", "Credentials, JWT issuance and customer profile"),
    ("Catalog", "Products, brands, categories", "Authoritative product metadata and price"),
    ("Cart", "Carts, cart lines", "Customer shopping intent"),
    ("Inventory", "Stock, reservations", "Atomic reserve/confirm/release/restock"),
    ("Order", "Orders, snapshots, outbox", "Checkout orchestration and durable order lifecycle"),
    ("Notification", "Notification log", "Consume confirmation events and record delivery"),
])

doc.add_heading("Technology baseline", level=2)
doc.add_paragraph("Java 21 · Spring Boot 3.3.5 · Spring Data JPA · Flyway · PostgreSQL · Spring Kafka · React · TypeScript · Vite · Docker Compose · JUnit/Mockito/Testcontainers")

page_break()
doc.add_heading("4. Critical checkout design", level=1)
doc.add_paragraph("Checkout is the system's highest-risk flow. The design deliberately stores each recoverable boundary before crossing to the next external operation.")
table(["Step", "Action", "Durable result / failure behavior"], [
    ("1", "Resolve customer-scoped idempotency key", "Existing key returns the original order"),
    ("2", "Load cart and snapshot catalog data", "Order keeps SKU, name and price history"),
    ("3", "Persist order and address", "PENDING exists before inventory call"),
    ("4", "Reserve inventory atomically", "Insufficient stock returns conflict; order becomes FAILED"),
    ("5", "Store reservation correlation", "Order becomes STOCK_RESERVED and is recoverable"),
    ("6", "Confirm inventory", "Idempotent inventory operation commits sold stock"),
    ("7", "Confirm order + insert outbox", "Both commit in one local database transaction"),
    ("8", "Clear cart", "Best-effort; never invalidates a confirmed order"),
    ("9", "Publish outbox event", "Worker retries Kafka with bounded exponential backoff"),
])

doc.add_heading("Order state model", level=2)
table(["State", "Meaning", "Allowed recovery"], [
    ("PENDING", "Order saved; inventory not durably correlated", "Reserve stock or fail"),
    ("STOCK_RESERVED", "Reservation id stored; confirmation may be interrupted", "Retry confirm through request or worker"),
    ("CONFIRMED", "Inventory sold and event durably queued", "Publish notification; customer may cancel"),
    ("FAILED", "Checkout did not cross a committed inventory boundary", "Start a new checkout"),
    ("CANCELLED", "Confirmed order cancelled and stock restored", "Terminal for MVP"),
])

doc.add_heading("Transactional outbox", level=2)
doc.add_paragraph("The order transaction writes both the CONFIRMED state and a unique ORDER_CONFIRMED outbox row. A scheduled publisher reads due PENDING rows, waits for Kafka acknowledgement, marks success as PUBLISHED, and records attempts, last error, and the next retry time on failure. The unique (order_id, event_type) constraint prevents duplicate enqueueing during reconciliation.")

page_break()
doc.add_heading("5. Data and API contracts", level=1)
doc.add_heading("Data ownership", level=2)
table(["Database", "Core tables"], [
    ("identity_db", "users, addresses"),
    ("catalog_db", "products, categories, brands"),
    ("cart_db", "carts, cart_items"),
    ("inventory_db", "inventory, reservations, reservation_items"),
    ("order_db", "orders, order_items, order_addresses, order_outbox"),
    ("notification_db", "notification_log"),
])
doc.add_paragraph("Services do not join across databases. Order items are immutable commercial snapshots so later catalog changes do not rewrite purchase history.")

doc.add_heading("External API groups", level=2)
table(["Route group", "Capability"], [
    ("/api/auth", "Register, login and token issuance"),
    ("/api/profile, /api/addresses", "Customer profile and address management"),
    ("/api/products, /api/categories, /api/brands", "Catalog discovery"),
    ("/api/cart", "Cart read and line mutations"),
    ("/api/inventory", "Availability plus internal reservation commands"),
    ("/api/orders", "Checkout, history, detail and cancellation"),
])

doc.add_heading("Event contract", level=2)
doc.add_paragraph("Topic: order.confirmed · Key: orderId · Payload: orderId, customerId, recipientEmail, itemCount, total, occurredAt. Delivery is at-least-once; the notification consumer uses a PostgreSQL partial unique index and INSERT ON CONFLICT DO NOTHING so each order creates one Kafka confirmation record.")

doc.add_heading("6. Security and operational posture", level=1)
for x in [
    "Passwords are hashed; identity issues signed JWTs and protected services verify them.",
    "Customer order reads are scoped by authenticated customer id; idempotency is also customer-scoped.",
    "Service health and OpenAPI endpoints support local diagnostics.",
    "Flyway owns schema evolution and Docker Compose provides a reproducible local environment.",
    "Current secrets are development configuration only and must move to managed secret storage before deployment.",
]: bullet(x)

page_break()
doc.add_heading("7. Verification and delivery status", level=1)
table(["Verification", "Result"], [
    ("Backend Maven reactor", "Passed"),
    ("GitHub Actions release gate", "Backend clean verify plus frontend npm ci, lint and build"),
    ("Order-service unit/integration tests", "7 passed; PostgreSQL Flyway V1–V4 applied"),
    ("Concurrent idempotency", "Verified with Testcontainers/PostgreSQL unique constraint"),
    ("Frontend production build", "Passed"),
    ("Docker service health", "Gateway and six services UP"),
    ("End-to-end checkout smoke test", "Confirmed order; inventory 100→98; duplicate key returned same order"),
    ("Kafka notification", "Consumed and logged as SENT; 8 concurrent duplicates create one row"),
])

doc.add_heading("Milestones", level=2)
table(["Milestone", "Status", "Outcome"], [
    ("0 — Foundation", "Complete", "Monorepo, gateway, databases, migrations, Compose"),
    ("1 — Catalog", "Complete", "Browse/search and product data"),
    ("2 — Identity", "Complete", "JWT authentication, profile and addresses"),
    ("3 — Cart", "Complete", "Persistent customer cart"),
    ("4 — Inventory", "Complete", "Atomic stock reservation lifecycle"),
    ("5 — Orders", "Complete", "Idempotent checkout, history and cancellation"),
    ("6 — Frontend journey", "Complete", "Integrated React customer experience"),
    ("7 — Release hardening", "In progress", "Outbox, consumer idempotency and CI complete"),
])

doc.add_heading("8. Next actions", level=1)
table(["Priority", "Action", "Exit criterion"], [
    ("P0", "Automate the full customer journey", "Register-to-cancellation smoke test runs repeatably"),
    ("P1", "Add tracing, metrics and alerts", "Checkout correlation visible across gateway, order, inventory and Kafka"),
    ("P1", "Add contract tests", "HTTP and event producer/consumer compatibility enforced"),
    ("P1", "Externalize secrets and environment config", "No deployment credentials stored in source"),
    ("P2", "Deploy shared staging environment", "Automated smoke test proves the complete customer journey"),
])
doc.add_paragraph("Recommendation: freeze new product features until the P0 items pass in CI. The MVP is now strong enough to demonstrate; the highest return comes from making that result repeatable and observable.")

# Footer and page numbers
for section in doc.sections:
    footer=section.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run=footer.add_run("OneStop MVP Architecture v0.2  •  "); run.font.size=Pt(8); run.font.color.rgb=RGBColor.from_string(GRAY)
    fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'), 'PAGE'); footer._p.append(fld)

doc.core_properties.title = "OneStop MVP Architecture v0.2"
doc.core_properties.subject = "Implemented MVP architecture and release-hardening plan"
doc.core_properties.author = "OneStop"
doc.save(OUT)
print(OUT)
